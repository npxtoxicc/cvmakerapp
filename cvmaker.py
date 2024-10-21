from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io

app = Flask(__name__)
app.secret_key = 'supersecretkey'

@app.route('/')
def index():
    return render_template('index.html')  # Dosya yolunu güncelledim

@app.route('/generate_cv', methods=['POST'])
def generate_cv():
    try:
        name = request.form['name']
        email = request.form['email']
        phone = request.form['phone']
        address = request.form['address']
        education = request.form['education']
        experience = request.form['experience']
        skills = request.form['skills']

        doc = Document()
        doc.add_heading('CV', 0)

        # Add personal details with a table for better alignment
        doc.add_heading('Personal Details', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = name
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'Email'
        hdr_cells[1].text = email
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = 'Phone'
        hdr_cells[1].text = phone
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = 'Address'
        hdr_cells[1].text = address

        # Add education with bullet points
        doc.add_heading('Education', level=1)
        education_paragraph = doc.add_paragraph()
        for edu in education.split('\n'):
            education_paragraph.add_run(f'• {edu}\n')

        # Add experience with bullet points
        doc.add_heading('Experience', level=1)
        experience_paragraph = doc.add_paragraph()
        for exp in experience.split('\n'):
            experience_paragraph.add_run(f'• {exp}\n')

        # Add skills with bullet points
        doc.add_heading('Skills', level=1)
        skills_paragraph = doc.add_paragraph()
        for skill in skills.split('\n'):
            skills_paragraph.add_run(f'• {skill}\n')

        # Save the document to a file
        file_path = f'{name}_CV.docx'
        doc.save(file_path)

        flash('CV successfully generated!', 'success')
        return send_file(file_path, as_attachment=True, download_name=f'{name}_CV.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        flash(f'An error occurred: {str(e)}', 'danger')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
